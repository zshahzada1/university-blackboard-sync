"""
FN585 Assignment 1 - DOCX Builder
Font: Arial 11pt | Line spacing: 1.5 | Per assessment brief
"""
import json
import os
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_LINE_SPACING, WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

os.chdir("/home/zo/University/FN585/Assignment1")

with open("outputs/results.json") as f:
    R = json.load(f)

def fmt(val, dp=4):
    return f"{val:.{dp}f}"

def pval_str(p):
    if p < 0.001: return "< 0.001***"
    elif p < 0.01: return f"{p:.4f}**"
    elif p < 0.05: return f"{p:.4f}*"
    else: return f"{p:.4f}"

def force_arial(run):
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "Arial")
    rFonts.set(qn("w:hAnsi"), "Arial")
    rFonts.set(qn("w:cs"), "Arial")
    rPr.insert(0, rFonts)

def set_font(run, size=11, bold=False, italic=False, code=False):
    run.font.name = "Courier New" if code else "Arial"
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if not code:
        force_arial(run)

def set_spacing(para, spacing=1.5, after=6):
    pf = para.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = spacing
    pf.space_after = Pt(after)

def add_para(doc, text, bold=False, size=11, spacing=1.5, after=6,
             align=WD_ALIGN_PARAGRAPH.LEFT, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, italic=italic)
    set_spacing(p, spacing, after)
    p.alignment = align
    return p

def add_heading(doc, num, title):
    p = doc.add_paragraph()
    run = p.add_run(f"{num}. {title}")
    set_font(run, size=11, bold=True)
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.5
    pf.space_before = Pt(12)
    pf.space_after = Pt(6)
    return p

def add_subheading(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, size=11, bold=True)
    set_spacing(p, 1.5, 4)
    return p

def add_figure(doc, path, caption, width=5.5):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(path, width=Inches(width))
    set_spacing(p, 1.0, 2)
    cap = doc.add_paragraph()
    run2 = cap.add_run(caption)
    set_font(run2, size=10, italic=True)
    set_spacing(cap, 1.0, 10)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

def shade_cell(cell, fill="D9D9D9"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)

def set_cell(row, idx, text, bold=False, size=10, align=WD_ALIGN_PARAGRAPH.LEFT, shade=None):
    cell = row.cells[idx]
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    set_font(run, size=size, bold=bold)
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(2)
    pf.space_after = Pt(2)
    if shade:
        shade_cell(cell, shade)

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# Title block
for text, size, bold in [
    ("FN585 Financial Modelling and Dealing", 14, True),
    ("Assignment 1: Time Series Analysis and Forecasting", 12, True),
    ("Zohaib Shahzada  |  FN585  |  2024-25", 11, False),
]:
    p = add_para(doc, text, bold=bold, size=size, spacing=1.0, after=3,
                 align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_paragraph()

# ==========================================================================
# Q1
# ==========================================================================
add_heading(doc, 1, "Annualised Inflation Rate")

q1_text = (
    "The annualised inflation rate is constructed as inflt = 400 * delta ln(PCECTPIt), "
    "converting quarterly log differences into an annual percentage. Figure 1 plots "
    "the resulting series from 1963:Q1 to 2013:Q4. The mean of inflation is clearly "
    "not constant over the sample. During the 1970s inflation rose sharply, peaking "
    "above 10% around 1974 and again near 1980, driven by oil price shocks and "
    "expansionary policy. After Volcker's disinflation in the early 1980s the series "
    "settled into a much lower range, averaging roughly 2-3% through the Great "
    "Moderation period. The variance also changes across sub-periods: the 1970s show "
    "large swings while the post-1990 era is comparatively calm. This shifting mean "
    "and heteroscedasticity are early signs that inflt may not be stationary."
)
add_para(doc, q1_text)
add_figure(doc, "figures/fig1_inflation.png",
    "Figure 1: U.S. annualised inflation rate, inflt = 400 x delta ln(PCECTPIt), 1963-2013.")

# ==========================================================================
# Q2
# ==========================================================================
add_heading(doc, 2, "Autocovariance and Autocorrelation Functions")

q2_text = (
    f"The sample autocovariance at lag 0 is {fmt(R['acov_lag0'], 4)}, which is just "
    f"the variance of inflt. Figure 2 plots the autocovariance function out to 20 lags, "
    f"and it declines only gradually. The ACF in Figure 3 tells the same story: the "
    f"lag-1 autocorrelation is {fmt(R['acf_lag1'], 4)} and each subsequent coefficient "
    f"remains well above zero. All 20 lags sit outside the 95% confidence bands. "
    f"This slow, near-linear decay is a textbook indicator of strong persistence. "
    f"When the ACF does not drop off quickly it suggests the series has a unit root "
    f"or is at least close to one, which fits with what the time-series plot already implied."
)
add_para(doc, q2_text)
add_figure(doc, "figures/fig2_autocovariance.png",
    "Figure 2: Sample autocovariance function of inflt, lags 0-20.")
add_figure(doc, "figures/fig3_acf_inflation.png",
    "Figure 3: Sample ACF of inflt, lags 1-20. Dashed lines show 95% confidence bands.")

# ==========================================================================
# Q3
# ==========================================================================
add_heading(doc, 3, "First Difference of Inflation Rate")

q3_text = (
    f"Taking the first difference of inflt removes much of the persistence visible in "
    f"the level series. Figure 4 plots delta inflt, which fluctuates around zero with "
    f"no obvious trend or wandering mean. The ACF in Figure 5 confirms the change: "
    f"the lag-1 autocorrelation drops to {fmt(R['diff_acf_lag1'], 4)}, which is negative "
    f"and small in absolute terms. Beyond lag 1 the autocorrelations are mostly inside "
    f"the 95% bands. This rapid decay contrasts sharply with the slow decline in the "
    f"level ACF. The pattern is consistent with delta inflt being I(0), meaning that "
    f"inflt itself is I(1) and needs one round of differencing to become stationary."
)
add_para(doc, q3_text)
add_figure(doc, "figures/fig4_diff_inflation.png",
    "Figure 4: First difference of inflation rate, delta inflt, 1963:Q2-2013:Q4.")
add_figure(doc, "figures/fig5_acf_diff_inflation.png",
    "Figure 5: ACF of delta inflt, lags 1-20.")

# ==========================================================================
# Q4
# ==========================================================================
add_heading(doc, 4, "Stationarity of inflt")

q4_text = (
    f"Three pieces of evidence bear on whether inflt is stationary. First, the time-series "
    f"plot in Figure 1 shows a mean that shifts between distinct regimes and volatility "
    f"that changes over time. A stationary process would not display these features. "
    f"Second, the ACF decays very slowly from {fmt(R['acf_lag1'], 4)} at lag 1, remaining "
    f"significant at all 20 lags, which is typical of a unit root process. Third, I run "
    f"the ADF test using ur.df() in R under three specifications: type=\"trend\" (which "
    f"includes a constant and time trend), type=\"drift\" (constant only), and type=\"none\" "
    f"(no deterministic terms). The drift specification is the most appropriate here because "
    f"inflt has a non-zero mean but no clear linear trend. Under drift the ADF test statistic "
    f"is {fmt(R['q4_stat'], 4)}, compared with critical values of {fmt(R['q4_cval_1pct'], 2)} "
    f"(1%), {fmt(R['q4_cval_5pct'], 2)} (5%), and {fmt(R['q4_cval_10pct'], 2)} (10%). The "
    f"test statistic is more negative than the 10% critical value but not more negative than "
    f"the 5% value, so we cannot reject the null of a unit root at the 5% significance level. "
    f"Taken together the visual, ACF, and formal test evidence all point to inflt being "
    f"non-stationary, or I(1)."
)
add_para(doc, q4_text)

# ==========================================================================
# Q5
# ==========================================================================
add_heading(doc, 5, "Stationarity of Delta inflt")

q5_text = (
    f"If inflt is I(1) then delta inflt should be I(0). The time-series plot of delta inflt "
    f"in Figure 4 oscillates around zero with no visible trend or drifting mean. The ACF "
    f"drops off immediately, with a lag-1 value of just {fmt(R['diff_acf_lag1'], 4)}. For "
    f"the ADF test I use type=\"none\" because once a series has been differenced there is "
    f"no reason to expect a non-zero intercept or deterministic trend. The test statistic is "
    f"{fmt(R['q5_stat'], 4)}, far more negative than the 1% critical value of "
    f"{fmt(R['q5_cval_1pct'], 2)} (5%: {fmt(R['q5_cval_5pct'], 2)}, 10%: "
    f"{fmt(R['q5_cval_10pct'], 2)}). We reject the null at the 1% level with no ambiguity. "
    f"Delta inflt is stationary, confirming that inflation is integrated of order one."
)
add_para(doc, q5_text)

# ==========================================================================
# Q6
# ==========================================================================
add_heading(doc, 6, "Autocorrelation of the Unemployment Rate")

q6_text = (
    f"Figure 6 shows the ACF of the unemployment rate. The lag-1 autocorrelation is "
    f"{fmt(R['urate_acf_lag1'], 4)}, extremely close to one, and all 20 lags remain "
    f"well outside the confidence bands. The decline is even slower than what we saw "
    f"for inflation in levels. This degree of persistence points toward either a unit "
    f"root or a highly persistent near-unit-root process in the unemployment rate."
)
add_para(doc, q6_text)
add_figure(doc, "figures/fig6_acf_urate.png",
    "Figure 6: ACF of uratet, 1962:Q3-2004:Q4, lags 1-20.")

# ==========================================================================
# Q7
# ==========================================================================
add_heading(doc, 7, "Stationarity of the Unemployment Rate")

q7_text = (
    f"The unemployment rate plot in Figure 7 shows cyclical fluctuations that recur "
    f"every few years but no sustained upward or downward trend, so the drift "
    f"specification is appropriate for the ADF test. The test statistic under drift is "
    f"{fmt(R['q7_stat'], 4)} against critical values of {fmt(R['q7_cval_1pct'], 2)} (1%), "
    f"{fmt(R['q7_cval_5pct'], 2)} (5%), and {fmt(R['q7_cval_10pct'], 2)} (10%). The "
    f"statistic is more negative than the 10% value but fails to beat the 5% critical "
    f"value, so we cannot reject the unit root null at conventional significance levels. "
    f"This result should be read carefully. Unemployment is bounded between 0 and 100 in "
    f"theory and between roughly 3 and 11 in practice, so a true unit root is implausible "
    f"in the long run. The series is more likely highly persistent with a near-unit root "
    f"in finite samples. For modelling purposes we treat it as non-stationary or "
    f"near-integrated over this sample."
)
add_para(doc, q7_text)
add_figure(doc, "figures/fig7_urate.png",
    "Figure 7: U.S. unemployment rate, uratet, 1962:Q3-2004:Q4.")

# ==========================================================================
# Q8
# ==========================================================================
add_heading(doc, 8, "Inflation Changes and Lagged Unemployment")

q8_text = (
    f"Figure 8 plots delta inflt against lagged unemployment, uratet-1. The OLS line "
    f"has a negative slope of {fmt(R['q8_slope'], 4)}, which aligns with the Phillips "
    f"curve intuition that higher unemployment is associated with falling inflation. "
    f"However, the relationship is weak. The HAC-robust p-value on the slope is "
    f"{fmt(R['q8_pval'], 4)}, only just below 5%, and the R-squared is {fmt(R['q8_rsq'], 4)}. "
    f"Less than 1% of the variation in inflation changes is explained by lagged "
    f"unemployment alone. The scatter is wide and noisy. A simple bivariate model is "
    f"clearly not enough here, which motivates including lagged values of both variables "
    f"in an ADL specification."
)
add_para(doc, q8_text)
add_figure(doc, "figures/fig8_scatter.png",
    f"Figure 8: delta inflt against lagged uratet-1, 1962:Q4-2004:Q4. "
    f"OLS slope = {fmt(R['q8_slope'],4)} (HAC p = {fmt(R['q8_pval'],4)}), R2 = {fmt(R['q8_rsq'],4)}.")

# ==========================================================================
# Q9
# ==========================================================================
add_heading(doc, 9, "ADL(2,2) Model: Estimation and Inference")

add_subheading(doc, "9(a)  Does lagged unemployment predict changes in inflation?")

q9a_text = (
    f"To test whether unemployment has predictive power for inflation changes beyond "
    f"what lagged delta inflt already captures, I test the joint null H0: delta1 = delta2 = 0. "
    f"The Wald F-statistic is {fmt(R['adl22_f_stat'], 4)} with a p-value of "
    f"{fmt(R['adl22_f_pval'], 4)}. We reject at the 1% level. Lagged unemployment "
    f"does carry information about future inflation movements that the autoregressive "
    f"terms alone miss. This supports the Phillips curve channel in the ADL framework."
)
add_para(doc, q9a_text)

# Table 1: ADL(2,2) results
add_para(doc,
    f"Table 1: ADL(2,2) estimation results. Dependent variable: delta inflt. "
    f"Sample: 1963:Q4-2004:Q4 (N = {int(R['adl22_nobs'])}). "
    f"HAC standard errors (Newey-West). R2 = {fmt(R['adl22_rsq'],4)}, "
    f"Adj. R2 = {fmt(R['adl22_adj_rsq'],4)}.",
    bold=True, size=10)

rows_data = [
    ("Constant",       "adl22_const",  "adl22_se_const",  "adl22_t_const",  "adl22_p_const"),
    ("delta inflt -1", "adl22_beta1",  "adl22_se_beta1",  "adl22_t_beta1",  "adl22_p_beta1"),
    ("delta inflt -2", "adl22_beta2",  "adl22_se_beta2",  "adl22_t_beta2",  "adl22_p_beta2"),
    ("uratet -1",      "adl22_delta1", "adl22_se_delta1", "adl22_t_delta1", "adl22_p_delta1"),
    ("uratet -2",      "adl22_delta2", "adl22_se_delta2", "adl22_t_delta2", "adl22_p_delta2"),
]

t1 = doc.add_table(rows=1 + len(rows_data), cols=5)
t1.style = "Table Grid"
hdrs = ["Variable", "Coefficient", "HAC Std. Error", "t-statistic", "p-value"]
for i, h in enumerate(hdrs):
    set_cell(t1.rows[0], i, h, bold=True, shade="D9D9D9")
for r_i, (label, c_key, se_key, t_key, p_key) in enumerate(rows_data):
    row = t1.rows[r_i + 1]
    set_cell(row, 0, label)
    set_cell(row, 1, fmt(R[c_key], 4), align=WD_ALIGN_PARAGRAPH.RIGHT)
    set_cell(row, 2, fmt(R[se_key], 4), align=WD_ALIGN_PARAGRAPH.RIGHT)
    set_cell(row, 3, fmt(R[t_key], 4), align=WD_ALIGN_PARAGRAPH.RIGHT)
    set_cell(row, 4, pval_str(R[p_key]), align=WD_ALIGN_PARAGRAPH.RIGHT)

add_para(doc,
    f"Note: *** p < 0.001, ** p < 0.01, * p < 0.05. "
    f"Joint F-test (H0: delta1 = delta2 = 0): F = {fmt(R['adl22_f_stat'],4)}, "
    f"p = {fmt(R['adl22_f_pval'],4)}. "
    f"Long-term effect of uratet on delta inflt: {fmt(R['adl22_lte'],4)}.",
    size=9, italic=True, spacing=1.0, after=10)

add_subheading(doc, "9(b)  Why are robust standard errors necessary?")

q9b_text = (
    "There are two reasons to use HAC standard errors rather than conventional OLS "
    "standard errors in this regression. The first is heteroscedasticity. Inflation "
    "volatility was substantially higher during the 1970s and early 1980s than in "
    "the post-1990 period, so the error variance is not constant across the sample. "
    "Ordinary OLS standard errors assume homoscedasticity and will be inconsistent "
    "when that assumption fails. The second problem is serial correlation. In a "
    "time-series regression with lagged dependent variables, the residuals are likely "
    "autocorrelated, which also invalidates conventional standard errors. The "
    "Newey-West estimator (Newey and West, 1987) corrects for both heteroscedasticity "
    "and autocorrelation simultaneously, producing standard errors that remain "
    "consistent under either or both violations."
)
add_para(doc, q9b_text)

add_subheading(doc, "9(c)  Are p = 2 and q = 2 lags appropriate?")

q9c_text = (
    f"Lag selection is based on the Bayesian Information Criterion. Table 2 reports "
    f"BIC values for all ADL(p,q) combinations with p and q ranging from 1 to 4. "
    f"The ADL(2,2) specification produces the lowest BIC at {fmt(R['bic_adl22'], 2)}, "
    f"confirming that two lags of each variable is the preferred choice. BIC is "
    f"better suited than AIC for this purpose because it imposes a heavier penalty "
    f"on additional parameters, which guards against over-fitting in finite samples. "
    f"The selection is data-driven rather than arbitrary, and the fact that BIC picks "
    f"a relatively parsimonious model is reassuring."
)
add_para(doc, q9c_text)

# Table 2: BIC grid
add_para(doc, "Table 2: BIC values across ADL(p,q) specifications.", bold=True, size=10)

bic_map = {
    (1,1):"bic_adl11",(1,2):"bic_adl12",(1,3):"bic_adl13",(1,4):"bic_adl14",
    (2,1):"bic_adl21",(2,2):"bic_adl22",(2,3):"bic_adl23",(2,4):"bic_adl24",
    (3,1):"bic_adl31",(3,2):"bic_adl32",(3,3):"bic_adl33",(3,4):"bic_adl34",
    (4,1):"bic_adl41",(4,2):"bic_adl42",(4,3):"bic_adl43",(4,4):"bic_adl44",
}
min_bic = min(R[v] for v in bic_map.values())

t2 = doc.add_table(rows=5, cols=5)
t2.style = "Table Grid"
set_cell(t2.rows[0], 0, "", shade="D9D9D9")
for j, q in enumerate([1,2,3,4]):
    set_cell(t2.rows[0], j+1, f"q = {q}", bold=True, shade="D9D9D9",
             align=WD_ALIGN_PARAGRAPH.CENTER)
for i, p in enumerate([1,2,3,4]):
    row = t2.rows[i+1]
    set_cell(row, 0, f"p = {p}", bold=True, shade="D9D9D9")
    for j, q in enumerate([1,2,3,4]):
        val = R[bic_map[(p,q)]]
        is_min = abs(val - min_bic) < 0.01
        sh = "C6EFCE" if is_min else None
        set_cell(row, j+1, fmt(val, 2), bold=is_min,
                 align=WD_ALIGN_PARAGRAPH.RIGHT, shade=sh)

add_para(doc,
    "Note: Green cell indicates the minimum BIC value across all 16 specifications.",
    size=9, italic=True, spacing=1.0, after=10)

add_subheading(doc, "9(d)  Why use delta inflt as dependent variable? Why the level of uratet?")

q9d_text = (
    "These two modelling choices rest on separate arguments. For the dependent variable, "
    "inflt is I(1) as established in Questions 4 and 5. Regressing one I(1) variable on "
    "another without cointegration produces a spurious regression where t-statistics and "
    "R-squared are unreliable (Granger and Newbold, 1974). First differencing transforms "
    "inflt into delta inflt, which is I(0), so standard inference applies. "
    "The treatment of unemployment is different. On theoretical grounds the expectations-augmented "
    "Phillips curve (Friedman, 1968; Phelps, 1968) says that the level of unemployment "
    "relative to the natural rate drives changes in inflation. Using the level of uratet "
    "rather than its difference preserves this economic relationship. On empirical grounds "
    "Stock and Watson (2015) note that unemployment in levels tends to give better forecast "
    "performance even when formal tests suggest it may be near-integrated. A differenced "
    "unemployment term would measure the acceleration of unemployment rather than its level, "
    "which is a different and less interpretable economic concept in the Phillips curve context."
)
add_para(doc, q9d_text)

# ==========================================================================
# Q10
# ==========================================================================
add_heading(doc, 10, "Out-of-Sample Forecast: Q1 2005")

# Table 3: lagged values
add_para(doc, "Table 3: Lagged values used in the Q1 2005 forecast.", bold=True, size=10)
t3 = doc.add_table(rows=5, cols=3)
t3.style = "Table Grid"
for i, h in enumerate(["Variable", "Period", "Value"]):
    set_cell(t3.rows[0], i, h, bold=True, shade="D9D9D9")
t3_data = [
    ("delta inflt -1", "2004:Q4", fmt(R["q10_dinfl_t1"], 4)),
    ("delta inflt -2", "2004:Q3", fmt(R["q10_dinfl_t2"], 4)),
    ("uratet -1",      "2004:Q4", fmt(R["q10_urate_t1"], 4)),
    ("uratet -2",      "2004:Q3", fmt(R["q10_urate_t2"], 4)),
]
for i, (var, per, val) in enumerate(t3_data):
    set_cell(t3.rows[i+1], 0, var)
    set_cell(t3.rows[i+1], 1, per, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell(t3.rows[i+1], 2, val, align=WD_ALIGN_PARAGRAPH.RIGHT)

doc.add_paragraph()

q10_text = (
    f"To forecast delta inflt for 2005:Q1 I plug the lagged values from Table 3 into "
    f"the estimated ADL(2,2) equation. The calculation proceeds as shown below. The "
    f"predicted change in inflation is {fmt(R['q10_forecast_diff'], 4)} percentage points. "
    f"Adding this to the 2004:Q4 inflation level of {fmt(R['q10_infl_2004q4'], 4)}% gives "
    f"a point forecast of inflt(2005:Q1) = {fmt(R['q10_forecast_infl'], 4)}%, a small decline "
    f"from the previous quarter. This is consistent with unemployment sitting above its "
    f"long-run average, which the model interprets as mild downward pressure on inflation."
)
add_para(doc, q10_text)

# Forecast equation
eq1 = (f"delta inflt(2005:Q1) = {fmt(R['adl22_const'],4)}"
       f" + ({fmt(R['adl22_beta1'],4)}) x {fmt(R['q10_dinfl_t1'],4)}"
       f" + ({fmt(R['adl22_beta2'],4)}) x {fmt(R['q10_dinfl_t2'],4)}"
       f" + ({fmt(R['adl22_delta1'],4)}) x {fmt(R['q10_urate_t1'],4)}"
       f" + ({fmt(R['adl22_delta2'],4)}) x {fmt(R['q10_urate_t2'],4)}"
       f" = {fmt(R['q10_forecast_diff'],4)}")
eq2 = (f"inflt(2005:Q1) = {fmt(R['q10_infl_2004q4'],4)}"
       f" + {fmt(R['q10_forecast_diff'],4)}"
       f" = {fmt(R['q10_forecast_infl'],4)}%")
for eq in [eq1, eq2]:
    p = add_para(doc, eq, size=10, italic=True, spacing=1.5, after=4,
                 align=WD_ALIGN_PARAGRAPH.CENTER)

# ==========================================================================
# References
# ==========================================================================
doc.add_page_break()
add_para(doc, "References", bold=True, size=12, after=10)
refs = [
    "Friedman, M. (1968) 'The role of monetary policy', American Economic Review, 58(1), pp. 1-17.",
    "Granger, C.W.J. and Newbold, P. (1974) 'Spurious regressions in econometrics', Journal of Econometrics, 2(2), pp. 111-120.",
    "Newey, W.K. and West, K.D. (1987) 'A simple, positive semi-definite, heteroskedasticity and autocorrelation consistent covariance matrix', Econometrica, 55(3), pp. 703-708.",
    "Phelps, E.S. (1968) 'Money-wage dynamics and labor-market equilibrium', Journal of Political Economy, 76(4), pp. 678-711.",
    "Phillips, A.W. (1958) 'The relation between unemployment and the rate of change of money wage rates in the United Kingdom, 1861-1957', Economica, 25(100), pp. 283-299.",
    "Stock, J.H. and Watson, M.W. (2015) Introduction to Econometrics. 3rd edn. Harlow: Pearson.",
]
for ref in refs:
    p = doc.add_paragraph()
    run = p.add_run(ref)
    set_font(run, size=11)
    set_spacing(p, 1.5, 4)
    pf = p.paragraph_format
    pf.first_line_indent = Inches(-0.3)
    pf.left_indent = Inches(0.3)

# ==========================================================================
# Appendix: R code
# ==========================================================================
doc.add_page_break()
add_para(doc, "Appendix: R Script", bold=True, size=12, after=6)
add_para(doc,
    "The script below was used to generate all figures, regression estimates, "
    "and forecast values in this report. Run from the Assignment1 directory with "
    "us_macro_quarterly.xlsx present.",
    size=11)

with open("assignment1.r", "r") as fh:
    rcode = fh.read()

code_para = doc.add_paragraph()
code_run = code_para.add_run(rcode)
code_run.font.name = "Courier New"
code_run.font.size = Pt(8)
pf = code_para.paragraph_format
pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
pf.line_spacing = 1.0
pf.space_after = Pt(0)

doc.save("Assignment1_FINAL.docx")
print("Assignment1_FINAL.docx saved.")
