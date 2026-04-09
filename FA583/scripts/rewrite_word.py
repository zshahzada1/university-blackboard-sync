from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
import copy

PATH = "/home/zo/University/FA583/Admin/RELX_FA583_FINAL_SUBMISSION.docx"
doc = Document(PATH)

# ─────────────────────────────────────────────────────────────
# 1. TABLE DATA
# ─────────────────────────────────────────────────────────────

PROF_HEADER = ["Ratio", "2022", "2023", "2024", "Formula"]
PROF_ROWS = [
    ["ROCE", "21.8%", "28.5%", "30.4%", "EBIT / (Total assets − Current liabilities)"],
    ["Gross profit margin", "64.4%", "64.9%", "65.0%", "Gross profit / Revenue"],
    ["Net profit margin", "19.1%", "19.4%", "20.5%", "Net profit (attrib. shareholders) / Revenue"],
]

LIQ_HEADER = ["Ratio", "2022", "2023", "2024", "Formula"]
LIQ_ROWS = [
    ["Current ratio", "0.59x", "0.52x", "0.52x", "Current assets / Current liabilities"],
    ["Quick assets ratio (acid test)", "0.53x", "0.46x", "0.47x", "(Current assets − Inventories) / Current liabilities"],
    ["Inventory holding period", "37.1 days", "36.1 days", "36.6 days", "(Inventories / Cost of sales) × 365"],
    ["Trade receivables collection period", "102.7 days", "92.6 days", "97.2 days", "(Trade receivables / Revenue) × 365"],
    ["Capital gearing ratio", "60.8%", "60.0%", "59.6%", "Long-term borrowings / (Long-term borrowings + Equity)"],
]


def set_table_data(table, header, data_rows):
    """Rebuild a table with the given header and data rows, preserving formatting where possible."""
    needed_rows = 1 + len(data_rows)  # header + data
    current_rows = len(table.rows)

    # Add rows if needed
    while len(table.rows) < needed_rows:
        table.add_row()

    # Remove extra rows (from the bottom)
    while len(table.rows) > needed_rows:
        tr = table.rows[-1]._tr
        table._tbl.remove(tr)

    all_rows = [header] + data_rows

    for r_idx, row_data in enumerate(all_rows):
        row = table.rows[r_idx]
        # Ensure enough columns
        while len(row.cells) < len(row_data):
            # Can't easily add cells; this shouldn't happen given the source tables have 5 cols
            pass

        for c_idx, text in enumerate(row_data):
            if c_idx >= len(row.cells):
                break
            cell = row.cells[c_idx]
            # Clear and rewrite
            for para in cell.paragraphs:
                for run in para.runs:
                    run.text = ""
            if cell.paragraphs:
                para = cell.paragraphs[0]
                para.clear()
                run = para.add_run(text)
                run.font.size = Pt(11)
                # Bold the header row
                if r_idx == 0:
                    run.font.bold = True
            else:
                para = cell.add_paragraph()
                run = para.add_run(text)
                run.font.size = Pt(11)
                if r_idx == 0:
                    run.font.bold = True


# ─────────────────────────────────────────────────────────────
# Identify the right tables
# Table 0 in doc: profitability (has Gross profit margin, Operating margin, etc.)
# Table 1 in doc: liquidity (has Current ratio, Quick ratio, etc.)
# Table 5 in doc: combined summary table (also has ratio headers — update this too)
# ─────────────────────────────────────────────────────────────

# Find profitability-only table: first table whose first non-header row is a ratio,
# and which does NOT contain 'LIQUIDITY' or 'SOLVENCY' section labels
# Also find the liquidity-only table.

# From inspection:
# Table 0 = profitability ratios only
# Table 1 = liquidity/solvency ratios only
# Table 5 = combined summary

def table_contains(tbl, text):
    for row in tbl.rows:
        for cell in row.cells:
            if text.lower() in cell.text.lower():
                return True
    return False

profitability_tables = []
liquidity_tables = []

for i, tbl in enumerate(doc.tables):
    has_ratio_header = tbl.rows and tbl.rows[0].cells[0].text.strip() == "Ratio"
    if not has_ratio_header:
        continue
    has_liq = table_contains(tbl, "Current ratio") or table_contains(tbl, "LIQUIDITY")
    has_prof = table_contains(tbl, "Gross profit") or table_contains(tbl, "PROFITABILITY") or table_contains(tbl, "ROCE")

    if has_prof and has_liq:
        # Combined table — update both sections
        profitability_tables.append((i, tbl, "combined"))
    elif has_prof:
        profitability_tables.append((i, tbl, "prof_only"))
    elif has_liq:
        liquidity_tables.append((i, tbl, "liq_only"))

print(f"Profitability tables: {[(x[0], x[2]) for x in profitability_tables]}")
print(f"Liquidity tables: {[(x[0], x[2]) for x in liquidity_tables]}")


def rebuild_combined_table(tbl):
    """Rebuild the combined summary table (Table 5): PROF section + LIQ section."""
    # Target structure:
    # Row 0: header
    # Row 1: PROFITABILITY (section label)
    # Rows 2-4: 3 profitability ratios
    # Row 5: LIQUIDITY & SOLVENCY (section label)
    # Rows 6-10: 5 liquidity/solvency ratios

    total_needed = 1 + 1 + len(PROF_ROWS) + 1 + len(LIQ_ROWS)
    # = 1 + 1 + 3 + 1 + 5 = 11

    while len(tbl.rows) < total_needed:
        tbl.add_row()
    while len(tbl.rows) > total_needed:
        tr = tbl.rows[-1]._tr
        tbl._tbl.remove(tr)

    all_rows_data = (
        [LIQ_HEADER]  # header (same for both)
        + [["PROFITABILITY", "", "", "", ""]]
        + PROF_ROWS
        + [["LIQUIDITY & SOLVENCY", "", "", "", ""]]
        + LIQ_ROWS
    )

    for r_idx, row_data in enumerate(all_rows_data):
        row = tbl.rows[r_idx]
        for c_idx, text in enumerate(row_data):
            if c_idx >= len(row.cells):
                break
            cell = row.cells[c_idx]
            for para in cell.paragraphs:
                para.clear()
            if cell.paragraphs:
                para = cell.paragraphs[0]
            else:
                para = cell.add_paragraph()
            run = para.add_run(text)
            run.font.size = Pt(11)
            # Bold header row and section label rows
            is_section = text in ("PROFITABILITY", "LIQUIDITY & SOLVENCY") and c_idx == 0
            if r_idx == 0 or is_section:
                run.font.bold = True


# Apply updates
for i, tbl, kind in profitability_tables:
    if kind == "prof_only":
        print(f"Updating profitability-only table {i}")
        set_table_data(tbl, PROF_HEADER, PROF_ROWS)
    elif kind == "combined":
        print(f"Updating combined table {i}")
        rebuild_combined_table(tbl)

for i, tbl, kind in liquidity_tables:
    if kind == "liq_only":
        print(f"Updating liquidity-only table {i}")
        # Liquidity-only table has section label rows ("LIQUIDITY", "SOLVENCY")
        # We replace the whole thing cleanly with 1 header + 5 data rows (no section labels)
        set_table_data(tbl, LIQ_HEADER, LIQ_ROWS)


# ─────────────────────────────────────────────────────────────
# 2. SECTION PROSE
# ─────────────────────────────────────────────────────────────

PROFITABILITY_TEXT = """RELX's ROCE rose from 21.8% in 2022 to 30.4% in 2024, an 8.6 percentage point improvement over two years. This reflects genuine operational progress rather than accounting effects. The Risk division grew faster than the group average and carries higher margins, so the revenue mix shift explains much of the gain. RELX's subscription model creates operating leverage: once content costs are covered, additional revenue flows through at high incremental margins.

Gross profit margin held at roughly 65% across all three years. For a business where most products are digital databases and analytics platforms, this stability is the point — content costs are scaling in line with revenue rather than running ahead of it.

Net profit margin improved from 19.1% to 20.5% despite the UK corporation tax rate rising from 19% to 25% in April 2023. A six percentage point tax increase and the margins still went up. That says something about how much operational headroom RELX has built into its cost structure."""

LIQUIDITY_TEXT = """The current ratio fell from 0.59x to 0.52x over the three years, and the quick assets ratio moved in parallel from 0.53x to 0.47x. Both sit below the conventional 1.0x benchmark. That looks like a problem until you examine the current liabilities figure. Around £4.1bn of it is deferred subscription income — cash RELX has already collected from clients for future periods. These liabilities get settled through content delivery, not cash payments. Strip that out and the picture changes considerably. The ratios reflect the business model, not financial stress.

Inventory holding has been consistent at 36–37 days throughout the period. For a business where most revenue comes from digital products, physical inventory is a relatively minor part of operations — mainly exhibition materials and print publications. The stability of this figure confirms there have been no supply-side disruptions worth noting.

Trade receivables collection improved from 102.7 days in 2022 to 97.2 days in 2024. The absolute figure looks long, but RELX's customers are large institutions and governments on annual contracts. Payment terms of 90+ days are standard in this market, and the direction of travel — a five-day improvement over two years — suggests tighter credit management.

The capital gearing ratio edged down from 60.8% to 59.6% between 2022 and 2024. The direction is right, but the absolute level stays elevated. RELX's equity base of £3.5bn sits under £8.2bn of goodwill from historical acquisitions. In a stress scenario, whether that equity could be realised at book value depends on acquisition valuations holding — which is not guaranteed."""

CONCLUSION_TEXT = """RELX's financial performance over 2022–2024 is strong on profitability and structurally complicated on liquidity. ROCE improved by nearly nine percentage points, margins held through a significant tax change, and revenue grew at a 5% compound rate.

The liquidity picture requires more careful reading. A current ratio of 0.52x and a quick ratio of 0.47x look alarming in isolation, but both are substantially explained by deferred subscription income in current liabilities. Capital gearing at around 60% with an intangible-heavy balance sheet is a more genuine concern.

Standard ratio analysis was not designed for subscription businesses. The receivables figure, the deferred income position, and the goodwill-dominated asset base all make the ratios harder to interpret than they would be for a manufacturer or retailer. Used alongside divisional organic revenue data and cash conversion figures, the ratios give a reasonable picture. Used in isolation, some would be misleading."""

SECTION_MAP = {
    "profitability": PROFITABILITY_TEXT,
    "liquidity": LIQUIDITY_TEXT,
    "conclusion": CONCLUSION_TEXT,
}


def is_section_heading(para, keyword):
    """Returns True if the paragraph looks like a section heading containing keyword."""
    text = para.text.strip().lower()
    return keyword in text and para.style.name.lower() in (
        "heading 1", "heading 2", "heading 3", "heading 4",
        "normal",  # in this doc headings use Normal style
    )


def looks_like_heading(para):
    """Heuristic: a paragraph is a heading if its text matches known section patterns."""
    text = para.text.strip().lower()
    # Known section starts in this document
    heading_patterns = [
        "1. company profile",
        "2. profitability",
        "3. liquidity",
        "4. conclusion",
        "references",
        "appendix",
        "table 1:",
        "table 2:",
        "source:",
    ]
    for pat in heading_patterns:
        if text.startswith(pat):
            return True
    return False


def replace_section_paragraphs(doc, section_keyword, new_text):
    """
    Find the heading paragraph for the section, then replace the following
    body paragraphs with new_text (split on blank lines into separate paragraphs).
    """
    paras = doc.paragraphs
    heading_idx = None

    for idx, para in enumerate(paras):
        if section_keyword in para.text.lower() and looks_like_heading(para):
            heading_idx = idx
            break

    if heading_idx is None:
        print(f"WARNING: Could not find heading for section '{section_keyword}'")
        return

    # Collect the body paragraph indices that follow the heading
    # Stop at the next heading or table caption
    body_indices = []
    for idx in range(heading_idx + 1, len(paras)):
        para = paras[idx]
        if looks_like_heading(para):
            break
        if para.text.strip() == "":
            continue  # skip blank paras but don't stop
        body_indices.append(idx)

    # Split new text into paragraphs (double newline = paragraph break)
    new_paras = [p.strip() for p in new_text.strip().split("\n\n") if p.strip()]

    print(f"Section '{section_keyword}': found {len(body_indices)} existing body paras, "
          f"need {len(new_paras)} new paras")

    # Replace existing paragraphs
    replaced = 0
    for i, body_idx in enumerate(body_indices):
        if i < len(new_paras):
            para = paras[body_idx]
            para.clear()
            run = para.add_run(new_paras[i])
            run.font.size = Pt(11)
            replaced += 1
        else:
            # More existing paras than new ones — clear the extras
            para = paras[body_idx]
            para.clear()

    # If we need more paragraphs than existed, insert after the last replaced one
    if len(new_paras) > replaced:
        # Find the actual paragraph element for the last replaced paragraph
        if body_indices:
            last_body_idx = body_indices[min(len(new_paras), len(body_indices)) - 1]
            last_para = paras[last_body_idx]
        else:
            last_para = paras[heading_idx]

        for extra_text in new_paras[replaced:]:
            new_p_elem = OxmlElement('w:p')
            last_para._element.addnext(new_p_elem)
            # Now we need to add a run inside it — create a temporary paragraph object
            # We do this by finding it in the refreshed paragraph list
            # Simpler: build the XML directly
            new_r = OxmlElement('w:r')
            new_rpr = OxmlElement('w:rPr')
            new_sz = OxmlElement('w:sz')
            new_sz.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', '22')  # 11pt = 22 half-pts
            new_szCs = OxmlElement('w:szCs')
            new_szCs.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', '22')
            new_rpr.append(new_sz)
            new_rpr.append(new_szCs)
            new_r.append(new_rpr)
            new_t = OxmlElement('w:t')
            new_t.text = extra_text
            new_t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
            new_r.append(new_t)
            new_p_elem.append(new_r)
            last_para = doc.paragraphs[doc.paragraphs.index(last_para) + 1] if False else last_para
            # Update last_para reference to the newly inserted paragraph
            # We need to find the new paragraph in doc.paragraphs — refresh the list
            # Actually use a simpler approach: track by element
            class FakePara:
                def __init__(self, elem):
                    self._element = elem
            last_para = FakePara(new_p_elem)


replace_section_paragraphs(doc, "profitability", PROFITABILITY_TEXT)
replace_section_paragraphs(doc, "liquidity", LIQUIDITY_TEXT)
replace_section_paragraphs(doc, "conclusion", CONCLUSION_TEXT)


# ─────────────────────────────────────────────────────────────
# 3. SAVE
# ─────────────────────────────────────────────────────────────

doc.save(PATH)
print("Done. Saved to", PATH)
