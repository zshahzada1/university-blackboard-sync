from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

PATH = "/home/zo/University/FA583/Admin/RELX_FA583_FINAL_SUBMISSION.docx"
doc = Document(PATH)

ALL_REFS = [
    "Elliott, B. and Elliott, J. (2019) Financial Accounting and Reporting. 19th edn. Harlow: Pearson Education Ltd.",
    "Fridson, M. and Alvarez, F. (2022) Financial Statement Analysis: A Practitioner's Guide. 5th edn. Hoboken, NJ: Wiley.",
    "Melville, A. (2019) International Financial Reporting: A practical guide. 7th edn. Harlow: Pearson Education Ltd.",
    "McKenzie, W. (2010) Financial Times Guide to Using and Interpreting Company Accounts. 4th edn. Harlow: FT Prentice Hall.",
    "Penman, S.H. (2013) Financial Statement Analysis and Security Valuation. 5th edn. New York: McGraw-Hill.",
    "RELX PLC (2022) Annual Report and Financial Statements 2022. London: RELX PLC.",
    "RELX PLC (2023) Annual Report and Financial Statements 2023. London: RELX PLC.",
    "RELX PLC (2024) Annual Report and Financial Statements 2024. London: RELX PLC.",
    "Thomson Reuters Corporation (2024) Annual Report 2024. Toronto: Thomson Reuters.",
    "Wolters Kluwer NV (2024) Annual Report 2024. Alphen aan den Rijn: Wolters Kluwer.",
]

# Check which references are already present
full = " ".join(p.text for p in doc.paragraphs)
missing = [r for r in ALL_REFS if r[:20] not in full]
print(f"Missing refs: {len(missing)}")
for r in missing:
    print(f"  - {r[:60]}")

# Find the References heading
paras = list(doc.paragraphs)
heading_idx = None
for i, para in enumerate(paras):
    if para.text.strip() == "References":
        heading_idx = i
        break

if heading_idx is not None:
    # Remove all existing ref paragraphs after the heading
    existing_ref_paras = []
    for p in paras[heading_idx + 1:]:
        if p.text.strip() and any(name in p.text for name in ["Elliott", "Fridson", "Melville", "McKenzie", "Penman", "RELX PLC", "Thomson Reuters", "Wolters Kluwer"]):
            existing_ref_paras.append(p)

    # Remove existing ref paragraphs
    for p in existing_ref_paras:
        p._element.getparent().remove(p._element)

    # Re-add all refs in order after the heading
    heading_para = paras[heading_idx]
    last = heading_para
    for ref_text in ALL_REFS:
        new_elem = OxmlElement("w:p")
        last._element.addnext(new_elem)
        new_para = Paragraph(new_elem, last._parent)
        run = new_para.add_run(ref_text)
        run.font.size = Pt(11)
        last = new_para
    print(f"Restored {len(ALL_REFS)} references after heading")
else:
    print("ERROR: cannot find References heading")

doc.save(PATH)
print("Saved")
