from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

PATH = "/home/zo/University/FA583/Admin/RELX_FA583_FINAL_SUBMISSION.docx"
doc = Document(PATH)

# ── Fix 2 & 3: paragraph text replacements ────────────────────────────────
OLD_PARALLEL = "Read in isolation the numbers can mislead \u2014 read alongside organic revenue growth and cash conversion data from the annual reports, they make considerably more sense."
NEW_PARALLEL_A = "Taken in isolation, the numbers can mislead."
NEW_PARALLEL_B = "The full picture needs annual report commentary on organic revenue growth and cash conversion alongside them."

OLD_OPENER = "RELX's 2022-2024 performance is strong where the numbers are easiest to read."
NEW_OPENER = "RELX's profitability numbers over 2022-2024 make a straightforward case."

for para in doc.paragraphs:
    if OLD_PARALLEL in para.text:
        # Replace the old sentence within the paragraph text, keeping the rest
        old_full = para.text
        new_full_a = old_full.replace(OLD_PARALLEL, NEW_PARALLEL_A)
        para.clear()
        run = para.add_run(new_full_a)
        run.font.size = Pt(11)
        # Insert second sentence as new paragraph after this one
        new_elem = OxmlElement("w:p")
        para._element.addnext(new_elem)
        new_para = Paragraph(new_elem, para._parent)
        run2 = new_para.add_run(NEW_PARALLEL_B)
        run2.font.size = Pt(11)
        print("Fixed: negative parallelism")
    elif OLD_OPENER in para.text:
        old_full = para.text
        new_full = old_full.replace(OLD_OPENER, NEW_OPENER)
        para.clear()
        run = para.add_run(new_full)
        run.font.size = Pt(11)
        print("Fixed: generic opener")

# ── Fix 1: Restore Melville reference if missing ──────────────────────────
full_text_after = " ".join(p.text for p in doc.paragraphs)
if "Melville" not in full_text_after:
    melville_text = "Melville, A. (2019) International Financial Reporting: A practical guide. 7th edn. Harlow: Pearson Education Ltd."
    # No References section exists — find the last non-empty paragraph and add References heading + Melville after it
    last_non_empty = None
    for para in doc.paragraphs:
        if para.text.strip():
            last_non_empty = para

    if last_non_empty is not None:
        # Add a blank line, then "References" heading, then Melville
        for label, text in [("heading", "References"), ("ref", melville_text)]:
            new_elem = OxmlElement("w:p")
            last_non_empty._element.addnext(new_elem)
            new_para = Paragraph(new_elem, last_non_empty._parent)
            run = new_para.add_run(text)
            run.font.size = Pt(11)
            if label == "heading":
                run.bold = True
            last_non_empty = new_para
        print("Restored: References section with Melville reference")
else:
    print("Melville already present — no action needed")

doc.save(PATH)
print("Saved")
