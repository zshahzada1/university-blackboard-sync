from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

PATH = "/home/zo/University/FA583/Admin/RELX_FA583_FINAL_SUBMISSION.docx"
doc = Document(PATH)

# ── Section prose replacements ────────────────────────────────────────────

SECTIONS = {
    "profitability": [
        "RELX's ROCE rose from 21.8% in 2022 to 30.4% in 2024. That 8.6 percentage point improvement is driven mainly by the Risk division, which grew faster than the rest of the group and runs at higher margins. The subscription model helps too: once content costs are covered, incremental revenue flows through at margins well above the headline figures.",
        "Gross profit margin held at roughly 65% across all three years. For a business where most products are digital databases and analytics platforms, this is the expected pattern. It means content costs are scaling with revenue, which is what you want to see.",
        "Net profit margin improved from 19.1% to 20.5% despite the UK corporation tax rate rising from 19% to 25% in April 2023. The tax change alone should have knocked about 1.5 points off the margin. That it did not reflects how much operating leverage RELX has built up over this period.",
    ],
    "liquidity": [
        "The current ratio fell from 0.59x to 0.52x over the three years, and the quick assets ratio tracked it closely, dropping from 0.53x to 0.47x. Both sit below the 1.0x level typically cited as adequate. What matters is what sits inside the current liabilities figure: around \u00a34.1bn of deferred subscription income, representing cash already collected for periods not yet delivered. That obligation gets settled with content, not cash outflows. Strip it out and both ratios look materially less concerning.",
        "Inventory holding period stayed around 36-37 days throughout. Most of RELX's revenue comes from digital products, so physical inventory is a small part of operations, mainly exhibition materials and some print. The number barely moved, which is about what you would expect from a business with this kind of product mix.",
        "Trade receivables collection came down from 102.7 days in 2022 to 97.2 days in 2024. On its own, 97 days looks long. RELX sells to governments, universities, and large enterprises on annual contracts, where 90-day terms are normal. The five-day improvement over two years points to tighter collection discipline.",
        "Capital gearing edged down from 60.8% to 59.6% between 2022 and 2024. That marginal improvement is better than nothing, but 60% remains high leverage, particularly when most of the equity supporting it is goodwill \u2014 \u00a38.2bn, accumulated from past acquisitions. Whether that goodwill holds its value under pressure is a question the ratios cannot answer.",
    ],
    "conclusion": [
        "RELX's 2022-2024 performance is strong where the numbers are easiest to read. ROCE rose from 21.8% to 30.4% and net profit margin improved despite a substantial tax increase. The profitability picture is clear.",
        "The liquidity ratios need more context. A current ratio of 0.52x and a quick ratio of 0.47x both look like problems until you account for the deferred subscription income sitting in current liabilities. Capital gearing at 60%, backed partly by goodwill, is the more substantive concern.",
        "Standard ratio analysis was built around tangible-asset businesses with clear distinctions between cash and non-cash obligations. RELX does not fit that model well. The deferred income complicates the liquidity picture; the goodwill complicates the gearing calculation. Read in isolation the numbers can mislead \u2014 read alongside organic revenue growth and cash conversion data from the annual reports, they make considerably more sense.",
    ],
}

def is_heading(para):
    style = para.style.name.lower()
    if "heading" in style:
        return True
    text = para.text.strip().lower()
    return any(text.startswith(f"{n}.") for n in range(1, 10))

def replace_section(doc, keyword, new_paras):
    paras = list(doc.paragraphs)
    in_section = False
    body = []
    for para in paras:
        if keyword in para.text.lower() and is_heading(para):
            in_section = True
            continue
        if in_section:
            if is_heading(para) and keyword not in para.text.lower():
                break
            body.append(para)

    if not body:
        print(f"WARNING: no body paras found under '{keyword}'")
        return

    for j, para in enumerate(body):
        if j < len(new_paras):
            para.clear()
            run = para.add_run(new_paras[j])
            run.font.size = Pt(11)
        else:
            para.clear()

    if len(new_paras) > len(body):
        last = body[-1]
        for extra in new_paras[len(body):]:
            new_elem = OxmlElement("w:p")
            last._element.addnext(new_elem)
            new_para = Paragraph(new_elem, last._parent)
            run = new_para.add_run(extra)
            run.font.size = Pt(11)
            last = new_para

    print(f"Section '{keyword}' updated")

for kw, paras in SECTIONS.items():
    replace_section(doc, kw, paras)

# ── Fix Table 6 (Appendix D) ──────────────────────────────────────────────

NEW_ROWS = [
    ["Category", "Ratio", "2022", "2023", "2024", "Formula"],
    ["Profitability", "ROCE", "21.8%", "28.5%", "30.4%", "EBIT / (Total assets \u2212 Current liabilities)"],
    ["Profitability", "Gross profit margin", "64.4%", "64.9%", "65.0%", "Gross profit / Revenue"],
    ["Profitability", "Net profit margin", "19.1%", "19.4%", "20.5%", "Net profit (attrib. shareholders) / Revenue"],
    ["Liquidity", "Current ratio", "0.59x", "0.52x", "0.52x", "Current assets / Current liabilities"],
    ["Liquidity", "Quick assets ratio (acid test)", "0.53x", "0.46x", "0.47x", "(Current assets \u2212 Inventories) / Current liabilities"],
    ["Liquidity", "Inventory holding period", "37.1 days", "36.1 days", "36.6 days", "(Inventories / Cost of sales) \u00d7 365"],
    ["Liquidity", "Trade receivables collection period", "102.7 days", "92.6 days", "97.2 days", "(Trade receivables / Revenue) \u00d7 365"],
    ["Liquidity", "Capital gearing ratio", "60.8%", "60.0%", "59.6%", "Long-term borrowings / (Long-term borrowings + Equity)"],
]

def make_tbl_xml(rows):
    tbl = OxmlElement("w:tbl")
    tblPr = OxmlElement("w:tblPr")
    tblStyle = OxmlElement("w:tblStyle")
    tblStyle.set(qn("w:val"), "TableGrid")
    tblPr.append(tblStyle)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), "0")
    tblW.set(qn("w:type"), "auto")
    tblPr.append(tblW)
    tbl.append(tblPr)
    for r_i, row_data in enumerate(rows):
        tr = OxmlElement("w:tr")
        for text in row_data:
            tc = OxmlElement("w:tc")
            tcPr = OxmlElement("w:tcPr")
            tcW = OxmlElement("w:tcW")
            tcW.set(qn("w:w"), "0")
            tcW.set(qn("w:type"), "auto")
            tcPr.append(tcW)
            tc.append(tcPr)
            p = OxmlElement("w:p")
            r = OxmlElement("w:r")
            rPr = OxmlElement("w:rPr")
            if r_i == 0:
                b = OxmlElement("w:b")
                rPr.append(b)
            r.append(rPr)
            t = OxmlElement("w:t")
            t.text = text
            t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
            r.append(t)
            p.append(r)
            tc.append(p)
            tr.append(tc)
        tbl.append(tr)
    return tbl

# Find the appendix combined table — has PROFITABILITY row with empty siblings
target = None
for tbl in doc.tables:
    all_text = " ".join(c.text for row in tbl.rows for c in row.cells)
    if "PROFITABILITY" in all_text and "ROCE" in all_text:
        target = tbl
        break

if target:
    new_tbl = make_tbl_xml(NEW_ROWS)
    target._tbl.getparent().replace(target._tbl, new_tbl)
    print("Table 6 rebuilt OK")
else:
    print("WARNING: Appendix D table not found")

doc.save(PATH)
print("Saved:", PATH)
