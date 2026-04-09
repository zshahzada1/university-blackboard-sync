from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

doc = Document("/home/zo/University/FA583/Admin/RELX_FA583_FINAL_SUBMISSION.docx")

# ── Read current state ──────────────────────────────────────────────────────
print("=== Current paragraphs (non-empty) ===")
for i, p in enumerate(doc.paragraphs):
    if p.text.strip():
        print(f"{i}: {p.text[:120]}")

# ── Paragraph replacements ──────────────────────────────────────────────────
replacements = [
    (
        "RELX PLC is a FTSE 100 information analytics group, incorporated",
        "RELX PLC is a FTSE 100 information analytics group, incorporated in England and listed in London, Amsterdam, and New York (RELX PLC, 2024). It runs four divisions: Risk (data analytics and decision tools, around 35% of 2024 revenue), Scientific, Technical and Medical (academic publishing and databases, roughly 30%), Legal (LexisNexis, approximately 26%), and Exhibitions (trade shows, 9%) (RELX PLC, 2024). About 75% of revenue comes through subscriptions, which makes the top line fairly predictable and gives management more forward visibility than most businesses of comparable size. Its main competitors are Thomson Reuters, Wolters Kluwer, and Clarivate.",
    ),
    (
        "RELX's ROCE rose from 21.8%",
        "RELX's ROCE rose from 21.8% in 2022 to 30.4% in 2024 (RELX PLC, 2022, 2024). That 8.6 percentage point improvement is driven mainly by the Risk division, which grew faster than the rest of the group and runs at higher margins. The subscription model helps too: once content costs are covered, incremental revenue flows through at margins well above the headline figures.",
    ),
    (
        "Gross profit margin held at roughly 65%",
        "Gross profit margin held at roughly 65% across all three years (RELX PLC, 2022, 2023, 2024). For a business where most products are digital databases and analytics platforms, this is the expected pattern. It means content costs are scaling with revenue, which is what you want to see.",
    ),
    (
        "Net profit margin improved from 19.1%",
        "Net profit margin improved from 19.1% to 20.5% despite the UK corporation tax rate rising from 19% to 25% in April 2023 (RELX PLC, 2022, 2024). The tax change alone should have knocked about 1.5 points off the margin. That it did not reflects how much operating leverage RELX has built up over this period.",
    ),
    (
        "The current ratio fell from 0.59x",
        "The current ratio fell from 0.59x to 0.52x over the three years, and the quick assets ratio tracked it closely, dropping from 0.53x to 0.47x (RELX PLC, 2022, 2024). Both sit below the 1.0x level typically cited as adequate (Melville, 2019). What matters is what sits inside the current liabilities figure: around £4.1bn of deferred subscription income, representing cash already collected for periods not yet delivered (RELX PLC, 2024). That obligation gets settled with content, not cash outflows. Strip it out and both ratios look materially less concerning.",
    ),
    (
        "Inventory holding period stayed around",
        "Inventory holding period stayed around 36-37 days throughout (RELX PLC, 2022, 2023, 2024). Most of RELX's revenue comes from digital products, so physical inventory is a small part of operations, mainly exhibition materials and some print. The number barely moved, which is about what you would expect from a business with this kind of product mix.",
    ),
    (
        "Trade receivables collection came down from 102.7",
        "Trade receivables collection came down from 102.7 days in 2022 to 97.2 days in 2024 (RELX PLC, 2022, 2024). On its own, 97 days looks long. RELX sells to governments, universities, and large enterprises on annual contracts, where 90-day terms are normal. The five-day improvement over two years points to tighter collection discipline.",
    ),
    (
        "Capital gearing edged down from 60.8%",
        "Capital gearing edged down from 60.8% to 59.6% between 2022 and 2024 (RELX PLC, 2022, 2024). That marginal improvement is better than nothing, but 60% remains high leverage, particularly when most of the equity supporting it is goodwill — £8.2bn, accumulated from past acquisitions (RELX PLC, 2024). Whether that goodwill holds its value under pressure is a question the ratios cannot answer.",
    ),
]

NEW_CONCLUSION = (
    "Standard ratio analysis was built around tangible-asset businesses with clear distinctions "
    "between cash and non-cash obligations. RELX does not fit that model well. The deferred income "
    "complicates the liquidity picture; the goodwill complicates the gearing calculation. Taken in "
    "isolation, the numbers can mislead. The ratios make more sense read alongside the company's own "
    "commentary on organic revenue growth and cash conversion \u2014 information that does not appear "
    "in the headline numbers (RELX PLC, 2022, 2023, 2024)."
)

CONCLUSION_TRIGGERS = ["tangible-asset", "Taken in isolation", "The full picture needs"]

# ── Apply main replacements ─────────────────────────────────────────────────
replaced = {k: False for k, _ in replacements}
for para in doc.paragraphs:
    txt = para.text
    for trigger, new_text in replacements:
        if trigger in txt and not replaced[trigger]:
            # Preserve the style but replace all runs
            style = para.style
            for run in para.runs:
                run.text = ""
            para.clear()
            run = para.add_run(new_text)
            run.font.size = Pt(11)
            para.style = style
            replaced[trigger] = True
            print(f"  REPLACED: {trigger[:60]}")
            break

for trigger, _ in replacements:
    if not replaced[trigger]:
        print(f"  WARNING – trigger not found: {trigger[:60]}")

# ── Conclusion consolidation ────────────────────────────────────────────────
conclusion_paras = []
for i, para in enumerate(doc.paragraphs):
    txt = para.text
    if any(t in txt for t in CONCLUSION_TRIGGERS):
        conclusion_paras.append((i, para))
        print(f"  CONCLUSION PARA found at index {i}: {txt[:80]}")

if conclusion_paras:
    # Write the full new conclusion into the first matching paragraph
    first_idx, first_para = conclusion_paras[0]
    style = first_para.style
    first_para.clear()
    run = first_para.add_run(NEW_CONCLUSION)
    run.font.size = Pt(11)
    first_para.style = style
    print(f"  CONCLUSION written to para index {first_idx}")

    # Delete any subsequent orphan paragraphs
    for _, orphan in conclusion_paras[1:]:
        p_elem = orphan._element
        p_elem.getparent().remove(p_elem)
        print("  ORPHAN conclusion paragraph deleted")
else:
    print("  WARNING – no conclusion trigger paragraphs found")

# ── Source notes after tables ───────────────────────────────────────────────
SOURCE_NOTE = "Source: Calculated from RELX PLC (2022, 2023, 2024). Ratio definitions: Melville (2019)."

def table_text(tbl):
    texts = []
    for row in tbl.rows:
        for cell in row.cells:
            texts.append(cell.text)
    return " ".join(texts)

def add_caption(tbl, text):
    new_p = OxmlElement("w:p")
    tbl._tbl.addnext(new_p)
    para = Paragraph(new_p, tbl._tbl.getparent())
    run = para.add_run(text)
    run.font.size = Pt(9)
    run.font.italic = True
    return para

notes_added = 0
for tbl in doc.tables:
    txt = table_text(tbl)
    is_profitability = "ROCE" in txt and "Gross profit margin" in txt and "Current ratio" not in txt
    is_liquidity = "Current ratio" in txt and "Capital gearing" in txt and "ROCE" not in txt
    if is_profitability:
        add_caption(tbl, SOURCE_NOTE)
        notes_added += 1
        print("  SOURCE NOTE added after Profitability table")
    elif is_liquidity:
        add_caption(tbl, SOURCE_NOTE)
        notes_added += 1
        print("  SOURCE NOTE added after Liquidity table")

print(f"  Total source notes added: {notes_added} (expect 2)")

# ── Save ────────────────────────────────────────────────────────────────────
doc.save("/home/zo/University/FA583/Admin/RELX_FA583_FINAL_SUBMISSION.docx")
print("\nDocument saved.")

# ── Verify ──────────────────────────────────────────────────────────────────
doc2 = Document("/home/zo/University/FA583/Admin/RELX_FA583_FINAL_SUBMISSION.docx")
full = " ".join(p.text for p in doc2.paragraphs)

checks = [
    ("(RELX PLC, 2024)", True),
    ("(RELX PLC, 2022, 2024)", True),
    ("(RELX PLC, 2022, 2023, 2024)", True),
    ("(Melville, 2019)", True),
    ("does not appear in the headline numbers (RELX PLC, 2022, 2023, 2024)", True),
    ("The full picture needs annual report commentary", False),
    ("Taken in isolation, the numbers can mislead.", True),
]
print("\n=== VERIFICATION ===")
for text, should_exist in checks:
    present = text in full
    status = "OK" if present == should_exist else "FAIL"
    print(f"[{status}] {'present' if should_exist else 'absent'}: {text[:80]}")

source_note = "Source: Calculated from RELX PLC (2022, 2023, 2024). Ratio definitions: Melville (2019)."
print(f"Source notes: {full.count(source_note)} (expect 2)")
